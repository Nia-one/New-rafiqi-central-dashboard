from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("artifacts/RafiQi_Data_Source_Process_Notes.docx")
NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F2F2F2"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, text, bold=False, color="000000", size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16 if level == 1 else 12)
    run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade(cell, NAVY)
        set_cell_text(cell, header, bold=True, color=WHITE, size=8)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            if r_idx % 2:
                shade(cells[idx], LIGHT_GREY)
            set_cell_text(cells[idx], value, bold=(idx == 0), size=7.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(9)

# Masthead
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("RAFIQI CENTRAL")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor.from_string(BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Data Source Process Notes")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Current live configuration | 31 July 2026")
r.font.name = "Calibri"
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string("666666")

callout = doc.add_table(rows=1, cols=1)
callout.style = "Table Grid"
shade(callout.cell(0, 0), LIGHT_BLUE)
set_cell_text(
    callout.cell(0, 0),
    "Important distinction: Enterprise is the demand/employer stream. FONO and Shram Park (SP) are supply/capacity channels. They connect operationally, but they are not the same process.",
    bold=True,
    color=NAVY,
    size=9,
)

add_heading(doc, "How to read this note", 2)
for text in [
    "User-input Google Sheet: team enters actual operational data here.",
    "Backend operating sheet: controlled/normalized tables read by the dashboard; users should not maintain these manually.",
    "Bot/backend: WhatsApp bot and sync services create or update normalized backend records.",
    "Hybrid: the page combines user-input Sheet data with bot/backend records.",
]:
    p = doc.add_paragraph(style=None)
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(1)
    p.add_run("• ").bold = True
    p.add_run(text)

add_heading(doc, "Process Note 1 — Page-wise source and flow", 1)
add_heading(doc, "Mode: Self Learn", 2)
self_learn_rows = [
    ("Overview", "Studios; TEAM_OCCUPANCY; TEAM_ENTERPRISE_OUTCOMES; TEAM_ESSENTIALS_SUMMARY", "Living_Hourly; Work_Hourly; Essentials_Hourly; Enterprise_Demand; Member_Activation; Dashboard_Content", "Combined operating summary from live inputs and normalized backend."),
    ("Living", "Studios (primary existing-studio occupancy); TEAM_OCCUPANCY (supplemental)", "Studio_Master; Living_Hourly; Enterprise_Demand; Allocation_Mismatch", "Existing Studios only. FONO/SP pipeline is excluded."),
    ("Work", "TEAM_ENTERPRISE_OUTCOMES where required", "Work_Hourly; Enterprise_Demand", "Bot/backend is primary; Sheet supplies enterprise outcome context."),
    ("Essentials", "Essentials; TEAM_ESSENTIALS_SUMMARY; TEAM_ESSENTIALS_BOT; TEAM_ESSENTIALS_INVENTORY", "Essentials_Hourly; Essentials_Cohorts; Essentials_Inventory", "User data and bot observations are normalized for dashboard use."),
    ("Member NPS", "TEAM_MEMBER_FEEDBACK", "Member_NPS_Dashboard; Member_NPS_Feedback; Member_NPS_Responses", "Feedback input is automatically derived and synced to backend NPS tables."),
    ("People", "TEAM_REQ_PEOPLE_ROSTER; TEAM_MEMBER_ACTIVATION", "People_Roster; People_Follow_Through; Member_Activation", "Roster and activation inputs sync to controlled people tables."),
    ("Learning history", "TEAM_LEARNING_HISTORY; action/evidence references", "Learning_History; Evidence_Log; Action_Log", "Learning observations and governed evidence form the history."),
]
add_table(doc, ["Page", "Google Sheet / input tabs", "Bot / backend tables", "Process"], self_learn_rows, [1.0, 2.35, 2.35, 1.7])

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "Process Note 1 — Page-wise source and flow", 1)
add_heading(doc, "Mode: Self Drive", 2)
self_drive_rows = [
    ("Cash & Control", "TEAM_FINANCE_DAILY; governance input tabs", "Finance_Daily; Policy_Registry; Approval_Log; Action_Log; Evidence_Log", "Finance inputs sync to backend; controls and evidence govern decisions."),
    ("Enterprise Demand", "TEAM_ENTERPRISE_OUTCOMES; TEAM_SHRAMPARK_DEMAND; Fono Funnel where matching context is needed", "Enterprise_Demand; Studio_Master; Living_Hourly; Member_Activation", "Tracks employer/member demand and matching; it is not the FONO/SP supply process."),
    ("Member Adds", "Fono Funnel; TEAM_MEMBER_ACTIVATION", "Living_Hourly; Member_Activation; Studio_Master; governance logs", "Shows contracted/onboarded FONO/SP potential versus occupied nests. Existing Studios are separate."),
    ("Member Engagement", "TEAM_MEMBER_FEEDBACK; TEAM_REQ_MEMBER_ENGAGEMENT; TEAM_LEARNING_HISTORY", "Member_NPS_Dashboard; Member_NPS_Feedback; Member_NPS_Responses; governance logs", "Feedback/engagement inputs are derived and synced to backend."),
    ("Member Savings", "Essentials; TEAM_ESSENTIALS_SUMMARY; TEAM_ESSENTIALS_BOT; TEAM_ESSENTIALS_INVENTORY", "Essentials_Hourly; Essentials_Inventory; Action_Log; Evidence_Log; Approval_Log", "Savings, margin and service evidence combine in a governed view."),
    ("Nia Margins", "TEAM_FINANCE_DAILY; TEAM_OCCUPANCY; TEAM_REQ_POLICY_REGISTRY; TEAM_REQ_ACTION_LOG; TEAM_REQ_EVIDENCE_LOG; TEAM_REQ_APPROVAL_LOG", "Finance_Daily; Living_Hourly; Policy_Registry; Action_Log; Evidence_Log; Approval_Log", "CM2, occupancy, policy control, owner, action and verification drive the verdict."),
    ("Nia Growth", "TEAM_OCCUPANCY; Fono Funnel; TEAM_SHRAMPARK_DEMAND; TEAM_REQ_SP_SUPPLY; TEAM_ENTERPRISE_OUTCOMES", "Studio_Master; Theatre_Master; Enterprise_Demand; Living_Hourly", "Combines existing occupancy with FONO/SP capacity and enterprise demand."),
    ("Despatch", "TEAM_REQ_INCIDENT_LOG; TEAM_REQ_ACTION_LOG; TEAM_REQ_EVIDENCE_LOG; TEAM_REQ_APPROVAL_LOG", "Incident_Log; Action_Log; Evidence_Log; Approval_Log; Allocation_Mismatch", "Bot or governed logs generate the operational exception queue."),
    ("Your Sign-Off", "TEAM_REQ_APPROVAL_LOG; TEAM_REQ_POLICY_REGISTRY; TEAM_REQ_EVIDENCE_LOG", "Approval_Log; Decision_Log; Policy_Registry; Evidence_Log", "Governed approval/evidence records determine sign-off status."),
]
add_table(doc, ["Page", "Google Sheet / input tabs", "Bot / backend tables", "Process"], self_drive_rows, [1.0, 2.35, 2.35, 1.7])

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "Process Note 2 — Google Sheet vs bot/backend", 1)
add_heading(doc, "Mode: Self Learn", 2)
classification_sl = [
    ("Overview", "Yes", "Yes", "Hybrid"),
    ("Living", "Yes — Studios is primary", "Yes", "Google Sheet direct + backend"),
    ("Work", "Limited/contextual", "Yes — primary", "Bot/backend primary"),
    ("Essentials", "Yes", "Yes", "Hybrid"),
    ("Member NPS", "Yes", "Yes — auto-derived sync", "Google Sheet → backend sync"),
    ("People", "Yes", "Yes — synced", "Google Sheet → backend sync"),
    ("Learning history", "Yes", "Yes — governed logs", "Google Sheet → backend + governance"),
]
add_table(doc, ["Page", "Google Sheet", "Bot/backend", "Primary model"], classification_sl, [1.45, 1.8, 2.0, 2.5])

add_heading(doc, "Mode: Self Drive", 2)
classification_sd = [
    ("Cash & Control", "Yes", "Yes", "Hybrid / Sheet → backend sync"),
    ("Enterprise Demand", "Yes", "Yes", "Hybrid"),
    ("Member Adds", "Yes — Fono Funnel primary", "Yes", "Google Sheet direct + backend"),
    ("Member Engagement", "Yes", "Yes — auto-derived sync", "Hybrid / Sheet → backend sync"),
    ("Member Savings", "Yes", "Yes", "Hybrid"),
    ("Nia Margins", "Yes", "Yes", "Hybrid / Sheet → backend sync"),
    ("Nia Growth", "Yes", "Yes", "Hybrid"),
    ("Despatch", "Yes — governed logs", "Yes — primary", "Bot/backend + governed Sheet logs"),
    ("Your Sign-Off", "Yes — governance input", "Yes — primary", "Backend/governance hybrid"),
]
add_table(doc, ["Page", "Google Sheet", "Bot/backend", "Primary model"], classification_sd, [1.45, 1.8, 2.0, 2.5])

add_heading(doc, "Operating rule", 2)
rule = doc.add_table(rows=1, cols=1)
rule.style = "Table Grid"
shade(rule.cell(0, 0), LIGHT_BLUE)
set_cell_text(
    rule.cell(0, 0),
    "Users should fill only the designated user-input tabs. Backend tabs are generated or synchronized by the dashboard services and bots; they should not be manually maintained unless a controlled exception is explicitly approved.",
    bold=True,
    color=NAVY,
    size=9,
)

for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("RafiQi Central | Data Source Process Notes")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("777777")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
